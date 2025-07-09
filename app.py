import os
import re
import shutil
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Inches
import pytesseract
import pandas as pd

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        files = request.files.getlist('pdf_files')
        output_choice = request.form.get('output_format')
        generate_word = output_choice in ['word', 'both']
        generate_excel = output_choice in ['excel', 'both']

        download_links = []

        for pdf_file in files:
            filename = secure_filename(pdf_file.filename)
            pdf_path = os.path.join(UPLOAD_FOLDER, filename)
            pdf_file.save(pdf_path)

            base_name = os.path.splitext(filename)[0]
            docx_file = os.path.join(OUTPUT_FOLDER, f"{base_name}_OCR_with_Images.docx")
            xlsx_file = os.path.join(OUTPUT_FOLDER, f"{base_name}_Extracted_Data.xlsx")

            pages = convert_from_path(pdf_path, dpi=200)
            all_extracted_data = []

            if generate_word:
                doc = Document()

            for i, img in enumerate(pages, start=1):
                img_path = os.path.join(UPLOAD_FOLDER, f"{base_name}_page_{i}.jpg")
                img.save(img_path)
                text = pytesseract.image_to_string(img, lang='eng')
                clean_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

                if generate_word:
                    doc.add_heading(f'Page {i}', level=1)
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph(clean_text)

                if generate_excel:
                    all_extracted_data.append({
                        'Page': i,
                        'Text Snippet': clean_text[:250].replace('\n', ' ')
                    })

                os.remove(img_path)

            if generate_word:
                doc.save(docx_file)
                download_links.append(('Word', docx_file))

            if generate_excel and all_extracted_data:
                df = pd.DataFrame(all_extracted_data)
                df.to_excel(xlsx_file, index=False)
                download_links.append(('Excel', xlsx_file))

            os.remove(pdf_path)

        return render_template('index.html', download_links=download_links)

    return render_template('index.html')

@app.route('/download/<path:filepath>')
def download_file(filepath):
    return send_file(filepath, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
